"""
MathDoc Studio v1.0 — AI 數學文件智慧轉換平台
支援 Word 內建方程式 (OMML) ∣ MathType 一鍵批次轉換

pip install streamlit python-docx latex2mathml lxml
streamlit run app.py
"""

import streamlit as st
import re
from io import BytesIO
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from docx.enum.text import WD_ALIGN_PARAGRAPH
from lxml import etree
import latex2mathml.converter

# ═══════════════════════════════════════════════════════════
#  自訂主題樣式
# ═══════════════════════════════════════════════════════════

CUSTOM_CSS = """
<style>
[data-testid="stSidebar"]{background:linear-gradient(180deg,#f0fdfa 0%,#fff 70%);}
.block-container{padding-top:1.2rem;}
.brand-box{background:linear-gradient(135deg,#0d9488,#2dd4bf,#a7f3d0);
  padding:1.5rem 2rem;border-radius:14px;color:#fff;text-align:center;
  margin-bottom:1.1rem;box-shadow:0 4px 16px rgba(13,148,136,.22);}
.brand-box h1{margin:0;font-size:1.85rem;font-weight:800;letter-spacing:.02em;}
.brand-box p{margin:.25rem 0 0;opacity:.9;font-size:.88rem;}
.stat-row{display:flex;gap:.7rem;margin:.6rem 0 .9rem;}
.stat-card{flex:1;background:#f0fdfa;border:1px solid #99f6e4;
  border-radius:10px;padding:.65rem .4rem;text-align:center;}
.stat-card .num{font-size:1.55rem;font-weight:700;color:#0d9488;}
.stat-card .lbl{font-size:.76rem;color:#6b7280;margin-top:.1rem;}
.step-label{display:inline-flex;align-items:center;gap:.4rem;
  font-size:1.05rem;font-weight:600;color:#1f2937;margin:.5rem 0 .2rem;}
.step-label .badge{display:inline-block;background:#0d9488;color:#fff;
  width:24px;height:24px;border-radius:50%;text-align:center;
  line-height:24px;font-size:.8rem;font-weight:700;}
div.stButton>button[kind="primary"]{
  background:linear-gradient(135deg,#0d9488,#14b8a6)!important;
  border:none!important;border-radius:10px!important;}
div.stButton>button[kind="primary"]:hover{
  background:linear-gradient(135deg,#0f766e,#0d9488)!important;}
div.stDownloadButton>button{border-radius:10px!important;
  border:2px solid #14b8a6!important;font-weight:600!important;}
div.stDownloadButton>button:hover{background:#f0fdfa!important;}
</style>
"""

# ═══════════════════════════════════════════════════════════
#  1. LaTeX → MathML → OMML 轉換引擎
# ═══════════════════════════════════════════════════════════

def latex_to_mathml(latex_str: str):
    try:
        return latex2mathml.converter.convert(latex_str)
    except Exception:
        return None


def _walk(node, parent):
    tag = etree.QName(node.tag).localname if "}" in str(node.tag) else str(node.tag)

    if tag == "math":
        for ch in node:
            _walk(ch, parent)
    elif tag == "mrow":
        for ch in node:
            _walk(ch, parent)
    elif tag == "mfrac":
        f = etree.SubElement(parent, qn("m:f"))
        num = etree.SubElement(f, qn("m:num"))
        den = etree.SubElement(f, qn("m:den"))
        children = list(node)
        if len(children) >= 2:
            _walk(children[0], num)
            _walk(children[1], den)
    elif tag == "msqrt":
        rad = etree.SubElement(parent, qn("m:rad"))
        pr = etree.SubElement(rad, qn("m:radPr"))
        dh = etree.SubElement(pr, qn("m:degHide"))
        dh.set(qn("m:val"), "1")
        etree.SubElement(rad, qn("m:deg"))
        e = etree.SubElement(rad, qn("m:e"))
        for ch in node:
            _walk(ch, e)
    elif tag == "mroot":
        rad = etree.SubElement(parent, qn("m:rad"))
        deg_el = etree.SubElement(rad, qn("m:deg"))
        e = etree.SubElement(rad, qn("m:e"))
        children = list(node)
        if len(children) >= 2:
            _walk(children[0], e)
            _walk(children[1], deg_el)
    elif tag == "msup":
        s = etree.SubElement(parent, qn("m:sSup"))
        e = etree.SubElement(s, qn("m:e"))
        sup = etree.SubElement(s, qn("m:sup"))
        children = list(node)
        if len(children) >= 2:
            _walk(children[0], e)
            _walk(children[1], sup)
    elif tag == "msub":
        s = etree.SubElement(parent, qn("m:sSub"))
        e = etree.SubElement(s, qn("m:e"))
        sub = etree.SubElement(s, qn("m:sub"))
        children = list(node)
        if len(children) >= 2:
            _walk(children[0], e)
            _walk(children[1], sub)
    elif tag == "msubsup":
        s = etree.SubElement(parent, qn("m:sSubSup"))
        e = etree.SubElement(s, qn("m:e"))
        sub = etree.SubElement(s, qn("m:sub"))
        sup = etree.SubElement(s, qn("m:sup"))
        children = list(node)
        if len(children) >= 3:
            _walk(children[0], e)
            _walk(children[1], sub)
            _walk(children[2], sup)
    elif tag == "munderover":
        s = etree.SubElement(parent, qn("m:sSubSup"))
        e = etree.SubElement(s, qn("m:e"))
        sub = etree.SubElement(s, qn("m:sub"))
        sup = etree.SubElement(s, qn("m:sup"))
        children = list(node)
        if len(children) >= 3:
            _walk(children[0], e)
            _walk(children[1], sub)
            _walk(children[2], sup)
    elif tag == "mover":
        acc = etree.SubElement(parent, qn("m:acc"))
        accPr = etree.SubElement(acc, qn("m:accPr"))
        children = list(node)
        if len(children) >= 2:
            chr_el = etree.SubElement(accPr, qn("m:chr"))
            chr_el.set(qn("m:val"), children[1].text or "")
        e = etree.SubElement(acc, qn("m:e"))
        if children:
            _walk(children[0], e)
    elif tag == "munder":
        children = list(node)
        if len(children) >= 2:
            s = etree.SubElement(parent, qn("m:sSub"))
            e = etree.SubElement(s, qn("m:e"))
            sub = etree.SubElement(s, qn("m:sub"))
            _walk(children[0], e)
            _walk(children[1], sub)
    elif tag == "mtable":
        mat = etree.SubElement(parent, qn("m:m"))
        for row in node:
            rt = etree.QName(row.tag).localname if "}" in str(row.tag) else str(row.tag)
            if rt == "mtr":
                mr = etree.SubElement(mat, qn("m:mr"))
                for cell in row:
                    ct = etree.QName(cell.tag).localname if "}" in str(cell.tag) else str(cell.tag)
                    if ct == "mtd":
                        e = etree.SubElement(mr, qn("m:e"))
                        for ch in cell:
                            _walk(ch, e)
    elif tag == "mfenced":
        d = etree.SubElement(parent, qn("m:d"))
        dPr = etree.SubElement(d, qn("m:dPr"))
        begChr = etree.SubElement(dPr, qn("m:begChr"))
        begChr.set(qn("m:val"), node.get("open", "("))
        endChr = etree.SubElement(dPr, qn("m:endChr"))
        endChr.set(qn("m:val"), node.get("close", ")"))
        e = etree.SubElement(d, qn("m:e"))
        for ch in node:
            _walk(ch, e)
    elif tag in ("mi", "mn", "mo", "mtext", "ms"):
        r = etree.SubElement(parent, qn("m:r"))
        if tag == "mi" and node.text and len(node.text) == 1:
            rPr = etree.SubElement(r, qn("m:rPr"))
            sty = etree.SubElement(rPr, qn("m:sty"))
            sty.set(qn("m:val"), "i")
        t = etree.SubElement(r, qn("m:t"))
        t.text = node.text or ""
    elif tag in ("mspace", "mpadded", "mstyle", "menclose"):
        for ch in node:
            _walk(ch, parent)
    else:
        if node.text and node.text.strip():
            r = etree.SubElement(parent, qn("m:r"))
            t = etree.SubElement(r, qn("m:t"))
            t.text = node.text
        for ch in node:
            _walk(ch, parent)


def _make_omml(latex_str):
    oMath = etree.Element(qn("m:oMath"))
    mathml_str = latex_to_mathml(latex_str)
    if mathml_str:
        try:
            tree = etree.fromstring(mathml_str.encode("utf-8"))
            _walk(tree, oMath)
            return oMath
        except Exception:
            pass
    r = etree.SubElement(oMath, qn("m:r"))
    t = etree.SubElement(r, qn("m:t"))
    t.text = latex_str
    return oMath


def _make_omml_display(latex_str):
    oMathPara = etree.Element(qn("m:oMathPara"))
    oMath = etree.SubElement(oMathPara, qn("m:oMath"))
    mathml_str = latex_to_mathml(latex_str)
    if mathml_str:
        try:
            tree = etree.fromstring(mathml_str.encode("utf-8"))
            _walk(tree, oMath)
            return oMathPara
        except Exception:
            pass
    r = etree.SubElement(oMath, qn("m:r"))
    t = etree.SubElement(r, qn("m:t"))
    t.text = latex_str
    return oMathPara


# ═══════════════════════════════════════════════════════════
#  2. Markdown 結構化解析器
# ═══════════════════════════════════════════════════════════

def parse_inline(text: str) -> list:
    parts = []
    pattern = r'(\$[^\$\n]+?\$|\*\*[^\*]+?\*\*)'
    last = 0
    for m in re.finditer(pattern, text):
        if m.start() > last:
            parts.append({"type": "text", "content": text[last:m.start()]})
        raw = m.group(0)
        if raw.startswith("$") and raw.endswith("$"):
            parts.append({"type": "math", "content": raw[1:-1].strip()})
        elif raw.startswith("**") and raw.endswith("**"):
            parts.append({"type": "bold", "content": raw[2:-2]})
        last = m.end()
    if last < len(text):
        parts.append({"type": "text", "content": text[last:]})
    return parts


def _is_bold_only(parts: list) -> bool:
    return all(
        p["type"] == "bold"
        or (p["type"] == "text" and not p["content"].strip())
        for p in parts
    )


def parse_document(text: str) -> list:
    text = text.replace("\\_", "_")
    lines = text.split("\n")
    segments = []
    i = 0

    def _is_table_start(idx):
        if idx >= len(lines) or idx + 1 >= len(lines):
            return False
        s = lines[idx].strip()
        if s.count("|") < 2:
            return False
        sep = lines[idx + 1].strip()
        return bool(re.match(r"^[\|\s:_-]+$", sep) and re.search(r"-{2,}", sep))

    def _is_special(idx):
        if idx >= len(lines):
            return True
        s = lines[idx].strip()
        if not s:
            return True
        if re.match(r"^-{3,}$", s) or re.match(r"^\*{3,}$", s):
            return True
        if re.match(r"^#{1,6}\s+", s):
            return True
        if s.startswith("$$"):
            return True
        if _is_table_start(idx):
            return True
        return False

    while i < len(lines):
        stripped = lines[i].strip()
        if not stripped:
            i += 1
            continue
        if re.match(r"^-{3,}$", stripped) or re.match(r"^\*{3,}$", stripped):
            segments.append({"type": "hr"})
            i += 1
            continue
        hm = re.match(r"^(#{1,6})\s+(.*)", stripped)
        if hm:
            segments.append({
                "type": "heading",
                "level": len(hm.group(1)),
                "content": hm.group(2).strip(),
            })
            i += 1
            continue
        if stripped.startswith("$$"):
            math_text = stripped[2:]
            if math_text.endswith("$$") and len(math_text) > 2:
                segments.append({"type": "display_math", "content": math_text[:-2].strip()})
                i += 1
            else:
                collect = [math_text] if math_text else []
                i += 1
                while i < len(lines):
                    if lines[i].strip().endswith("$$"):
                        tail = lines[i].strip()[:-2]
                        if tail:
                            collect.append(tail)
                        i += 1
                        break
                    collect.append(lines[i])
                    i += 1
                segments.append({"type": "display_math", "content": "\n".join(collect).strip()})
            continue
        if _is_table_start(i):
            headers = [c.strip() for c in stripped.strip("|").split("|")]
            i += 2
            rows = []
            while i < len(lines) and "|" in lines[i] and lines[i].strip():
                row = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                rows.append(row)
                i += 1
            segments.append({"type": "table", "headers": headers, "rows": rows})
            continue
        para_lines = [stripped]
        i += 1
        while i < len(lines) and not _is_special(i):
            para_lines.append(lines[i].strip())
            i += 1
        full_text = " ".join(para_lines)
        segments.append({"type": "paragraph", "parts": parse_inline(full_text)})

    merged = []
    j = 0
    while j < len(segments):
        seg = segments[j]
        if (seg["type"] == "paragraph"
                and _is_bold_only(seg["parts"])
                and j + 1 < len(segments)
                and segments[j + 1]["type"] == "paragraph"):
            new_parts = (seg["parts"]
                         + [{"type": "text", "content": " "}]
                         + segments[j + 1]["parts"])
            merged.append({"type": "paragraph", "parts": new_parts})
            j += 2
        else:
            merged.append(seg)
            j += 1
    return merged


# ═══════════════════════════════════════════════════════════
#  3. 內容分析引擎（新增）
# ═══════════════════════════════════════════════════════════

def analyze_segments(segments: list) -> dict:
    stats = {"inline": 0, "display": 0, "tables": 0,
             "paragraphs": 0, "headings": 0}
    for seg in segments:
        t = seg["type"]
        if t == "display_math":
            stats["display"] += 1
        elif t == "heading":
            stats["headings"] += 1
        elif t == "table":
            stats["tables"] += 1
            for row in seg.get("rows", []):
                for cell in row:
                    for p in parse_inline(cell):
                        if p["type"] == "math":
                            stats["inline"] += 1
        elif t == "paragraph":
            stats["paragraphs"] += 1
            for p in seg.get("parts", []):
                if p["type"] == "math":
                    stats["inline"] += 1
    stats["total_eq"] = stats["inline"] + stats["display"]
    return stats


def extract_equations(segments: list) -> list:
    eqs = []
    for seg in segments:
        if seg["type"] == "display_math":
            eqs.append(("展示", seg["content"]))
        elif seg["type"] == "paragraph":
            for p in seg.get("parts", []):
                if p["type"] == "math":
                    eqs.append(("行內", p["content"]))
        elif seg["type"] == "table":
            for row in seg.get("rows", []):
                for cell in row:
                    for p in parse_inline(cell):
                        if p["type"] == "math":
                            eqs.append(("表格", p["content"]))
    return eqs


# ═══════════════════════════════════════════════════════════
#  4. Word 文件生成（增強：可自訂字型/字級/行距/表格色）
# ═══════════════════════════════════════════════════════════

LATEX_COLOR = RGBColor(0x7B, 0x2D, 0x8B)

TEMPLATES = {
    "📄 學術論文": {"font": "Times New Roman", "size": 12, "spacing": 1.5,  "tcolor": "D9E2F3"},
    "📝 課堂作業": {"font": "Calibri",         "size": 11, "spacing": 1.15, "tcolor": "E2EFDA"},
    "📋 考卷試題": {"font": "Times New Roman", "size": 12, "spacing": 1.0,  "tcolor": "F2F2F2"},
    "⚙️ 自訂設定": None,
}

FONT_OPTIONS = ["Times New Roman", "Calibri", "Arial", "Cambria", "Georgia"]
SIZE_OPTIONS = [10, 11, 12, 14]
SPACING_OPTIONS = [1.0, 1.15, 1.5, 2.0]
COLOR_MAP = {
    "淡藍": "D9E2F3", "淡綠": "E2EFDA", "淡灰": "F2F2F2",
    "淡橘": "FCE4D6", "淡紫": "E8DAEF", "無底色": "FFFFFF",
}


def _add_inline_parts(para, parts, font_name, font_size, math_mode):
    for part in parts:
        if part["type"] == "text":
            run = para.add_run(part["content"])
            run.font.name = font_name
            run.font.size = font_size
        elif part["type"] == "bold":
            run = para.add_run(part["content"])
            run.bold = True
            run.font.name = font_name
            run.font.size = font_size
        elif part["type"] == "math":
            if math_mode == "latex":
                run = para.add_run("$" + part["content"] + "$")
                run.font.name = "Cambria Math"
                run.font.size = font_size
                run.font.color.rgb = LATEX_COLOR
            else:
                omml = _make_omml(part["content"])
                para._element.append(omml)


def _add_table(doc, headers, rows, math_mode, font_name, font_size, tcolor):
    ncols = len(headers)
    nrows = len(rows)
    table = doc.add_table(rows=1 + nrows, cols=ncols, style="Table Grid")
    table.autofit = False
    total_cm = 16.0
    if ncols == 2:
        widths = [3.0, total_cm - 3.0]
    elif ncols == 3:
        widths = [1.5, 2.0, total_cm - 3.5]
    elif ncols == 4:
        widths = [1.5, 2.0, 2.5, total_cm - 6.0]
    else:
        widths = [total_cm / ncols] * ncols
    for j in range(ncols):
        for row_obj in table.rows:
            row_obj.cells[j].width = Cm(widths[j])
    cell_fs = Pt(max(font_size.pt - 1, 9))
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ""
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_before = Pt(2)
        para.paragraph_format.space_after = Pt(2)
        run = para.add_run(h.replace("**", ""))
        run.bold = True
        run.font.name = font_name
        run.font.size = cell_fs
        try:
            shading = parse_xml(
                f'<w:shd {nsdecls("w")} w:fill="{tcolor}" w:val="clear"/>'
            )
            cell._element.get_or_add_tcPr().append(shading)
        except Exception:
            pass
    for i, row in enumerate(rows):
        for j in range(min(len(row), ncols)):
            cell = table.rows[i + 1].cells[j]
            cell.text = ""
            para = cell.paragraphs[0]
            para.paragraph_format.space_before = Pt(1)
            para.paragraph_format.space_after = Pt(1)
            parts = parse_inline(row[j])
            _add_inline_parts(para, parts, font_name, cell_fs, math_mode)
    doc.add_paragraph("")


def build_docx(segments, math_mode="omml", font_name="Times New Roman",
               body_size=12, line_spacing=1.15, tcolor="D9E2F3"):
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = font_name
    style.font.size = Pt(body_size)
    style.paragraph_format.space_after = Pt(2)
    style.paragraph_format.space_before = Pt(2)
    style.paragraph_format.line_spacing = line_spacing
    fs = Pt(body_size)

    for seg in segments:
        t = seg["type"]
        if t == "heading":
            level = min(seg["level"], 4)
            text = seg["content"].replace("**", "")
            if level <= 2:
                doc.add_heading(text, level=level)
            else:
                para = doc.add_paragraph()
                run = para.add_run(text)
                run.bold = True
                run.font.name = font_name
                run.font.size = Pt(14) if level == 3 else fs
                para.paragraph_format.space_before = Pt(8)
                para.paragraph_format.space_after = Pt(3)
        elif t == "hr":
            para = doc.add_paragraph()
            pPr = para._element.get_or_add_pPr()
            pBdr = parse_xml(
                f'<w:pBdr {nsdecls("w")}>'
                '<w:bottom w:val="single" w:sz="6" w:space="1" w:color="999999"/>'
                "</w:pBdr>"
            )
            pPr.append(pBdr)
        elif t == "display_math":
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if math_mode == "latex":
                run = para.add_run("$$" + seg["content"] + "$$")
                run.font.name = "Cambria Math"
                run.font.size = fs
                run.font.color.rgb = LATEX_COLOR
            else:
                omml = _make_omml_display(seg["content"])
                para._element.append(omml)
        elif t == "table":
            _add_table(doc, seg["headers"], seg["rows"],
                       math_mode, font_name, fs, tcolor)
        elif t == "paragraph":
            para = doc.add_paragraph()
            _add_inline_parts(para, seg["parts"], font_name, fs, math_mode)
    return doc


# ═══════════════════════════════════════════════════════════
#  5. MathML HTML 匯出
# ═══════════════════════════════════════════════════════════

def build_mathml_html(segments: list) -> str:
    parts = [
        "<!DOCTYPE html>",
        '<html><head><meta charset="utf-8"><title>MathML Export</title></head>',
        '<body style="font-family:Times New Roman;font-size:14pt;'
        'max-width:900px;margin:auto;padding:20px">',
        "<h2>方程式（MathML 格式）</h2><hr>",
    ]
    idx = 1
    for seg in segments:
        if seg["type"] == "display_math":
            mathml = latex_to_mathml(seg["content"])
            if mathml:
                parts.append(f"<p><b>#{idx}</b> <code>{seg['content']}</code></p>")
                parts.append(f'<div style="text-align:center;margin:1em 0">{mathml}</div>')
                idx += 1
        elif seg["type"] in ("paragraph", "table"):
            cells = []
            if seg["type"] == "paragraph":
                cells = [seg["parts"]]
            else:
                for row in seg["rows"]:
                    for cell in row:
                        cells.append(parse_inline(cell))
            for cell_parts in cells:
                for p in cell_parts:
                    if p["type"] == "math":
                        mathml = latex_to_mathml(p["content"])
                        if mathml:
                            parts.append(f"<p><b>#{idx}</b> <code>{p['content']}</code></p>")
                            parts.append(f"<p>{mathml}</p>")
                            idx += 1
    parts.append("</body></html>")
    return "\n".join(parts)


# ═══════════════════════════════════════════════════════════
#  6. 常數：Prompt / 說明 / VBA
# ═══════════════════════════════════════════════════════════

AI_PROMPT = (
    "請將所有產出的內容放入單一個『程式碼區塊（Code block）』中輸出，"
    "並嚴格遵守以下格式，以利後續程式解析：\n"
    "1. 排版保留：請使用標準 Markdown 語法處理文字排版（如 ### 標題、** 粗體）。\n"
    "2. 表格規範：若有表格，請使用標準 Markdown 表格（| 表格 |），"
    "並且【表格的正上方與正下方，務必各保留一個空白行】。\n"
    "3. 數學公式：請一律使用標準 LaTeX 語法呈現"
    "（行內公式使用 $ $，獨立行公式使用 $$ $$）。"
)

MATHTYPE_GUIDE = """
#### 🔧 MathType 批次轉換（僅需操作一次）

1. 用 Word 開啟下載的 `.docx` 檔案
2. 點選上方 **MathType** 索引標籤 → **Convert Equations**
3. 設定如下：
   - **Equation types to convert：** 勾選 ☑ **MathType translator text equations**
   - **Convert equations to：** MathType equations (OLE objects)
   - **Translator：** 選擇 **TeX -- AMS-LaTeX**
4. 按下 **Convert** → ✅ 完成

> `$...$`（行內）和 `$$...$$`（展示）是 MathType 預設分隔符，**無需額外設定**。
>
> 也可選取單一公式後按 **Alt+\\\\** 即時轉換。
"""

VBA_OMML = r'''
' ★ 推薦做法（免 VBA）
' MathType → Convert Equations → Whole Document
'   → MathType equations (OLE) → ☑ Word OMML equations → Convert

Sub ConvertOMMLtoMathType()
    On Error GoTo ErrHandler
    Dim total As Long
    total = ActiveDocument.OMaths.Count
    If total = 0 Then
        MsgBox "文件中未偵測到 OMML 方程式。", vbInformation
        Exit Sub
    End If
    Dim i As Long
    For i = total To 1 Step -1
        ActiveDocument.OMaths(i).Range.Select
        Selection.Copy
        Selection.Delete
        Selection.PasteSpecial Link:=False, DataType:=wdPasteOLEObject, Placement:=wdInLine
        DoEvents
    Next i
    MsgBox "已成功轉換 " & total & " 個方程式！", vbInformation
    Exit Sub
ErrHandler:
    MsgBox "轉換失敗，請確認 MathType 已安裝。", vbExclamation
End Sub
'''


# ═══════════════════════════════════════════════════════════
#  7. Streamlit 主介面
# ═══════════════════════════════════════════════════════════

def main():
    st.set_page_config(
        page_title="MathDoc Studio",
        page_icon="∫",
        layout="wide",
        initial_sidebar_state="expanded",
    )
    st.markdown(CUSTOM_CSS, unsafe_allow_html=True)

    # ── 初始化 session state ──
    for key, default in [("user_text", ""), ("show_analysis", False),
                          ("_uploaded_name", None)]:
        if key not in st.session_state:
            st.session_state[key] = default

    # ═══════════════════════════════════════
    #  側邊欄：設定面板
    # ═══════════════════════════════════════
    with st.sidebar:
        st.markdown("## ∫ MathDoc Studio")
        st.caption("v1.0 · AI 數學文件轉換平台")
        st.divider()

        # 輸出模式
        st.markdown("##### 📤 輸出模式")
        mode = st.radio(
            "mode", label_visibility="collapsed",
            options=[
                "Word 內建方程式 (OMML)",
                "MathType 模式 ⭐推薦",
                "MathType 傳統 (OMML+MathML)",
            ],
            index=1,
        )
        st.divider()

        # 文件樣式
        st.markdown("##### 🎨 文件樣式")
        tpl_name = st.selectbox(
            "範本", list(TEMPLATES.keys()), index=0,
            label_visibility="collapsed",
        )
        tpl = TEMPLATES[tpl_name]

        if tpl is not None:
            sel_font = tpl["font"]
            sel_size = tpl["size"]
            sel_spacing = tpl["spacing"]
            sel_tcolor_hex = tpl["tcolor"]
        else:
            sel_font = st.selectbox("字型", FONT_OPTIONS, index=0)
            sel_size = st.select_slider("字體大小 (pt)", SIZE_OPTIONS, value=12)
            sel_spacing = st.select_slider("行距", SPACING_OPTIONS, value=1.15)
            tc_name = st.selectbox("表格標題色", list(COLOR_MAP.keys()), index=0)
            sel_tcolor_hex = COLOR_MAP[tc_name]

        if tpl is not None:
            st.caption(f"字型 `{sel_font}` · {sel_size}pt · 行距 {sel_spacing}")

        st.divider()

        # AI Prompt
        with st.expander("🤖 AI Prompt 範本"):
            st.code(AI_PROMPT, language=None)
            st.caption("貼在 ChatGPT / Claude / Gemini 問題的最後面。")

        # MathType 步驟
        if "MathType" in mode:
            with st.expander("📋 MathType 轉換步驟"):
                st.markdown(MATHTYPE_GUIDE)

    # ═══════════════════════════════════════
    #  主區域
    # ═══════════════════════════════════════
    st.markdown(
        '<div class="brand-box">'
        "<h1>∫ MathDoc Studio</h1>"
        "<p>將 AI 生成的數學內容一鍵轉換為專業 Word 文件</p>"
        "</div>",
        unsafe_allow_html=True,
    )

    # ── Step 1：輸入 ──
    st.markdown(
        '<div class="step-label"><span class="badge">1</span>貼上或上傳內容</div>',
        unsafe_allow_html=True,
    )

    col_up, col_clr = st.columns([9, 1])
    with col_up:
        uploaded = st.file_uploader(
            "上傳 .txt / .md", type=["txt", "md"],
            label_visibility="collapsed",
        )
    with col_clr:
        if st.button("🗑️", help="清空輸入"):
            st.session_state["user_text"] = ""
            st.session_state["show_analysis"] = False
            st.session_state["_uploaded_name"] = None
            st.rerun()

    if uploaded is not None:
        if st.session_state.get("_uploaded_name") != uploaded.name:
            st.session_state["user_text"] = uploaded.read().decode("utf-8")
            st.session_state["_uploaded_name"] = uploaded.name
            st.rerun()

    text = st.text_area(
        "輸入區", height=320, label_visibility="collapsed",
        key="user_text",
        placeholder="在此貼上 AI 產生的 Markdown + LaTeX 內容（Ctrl+V）…",
    )

    # ── Step 2：分析 & 預覽 ──
    if text.strip():
        if st.button("🔍 分析內容", use_container_width=True):
            st.session_state["show_analysis"] = True

    if st.session_state["show_analysis"] and text.strip():
        segments = parse_document(text)
        stats = analyze_segments(segments)
        equations = extract_equations(segments)

        st.markdown(
            '<div class="step-label"><span class="badge">2</span>分析結果</div>',
            unsafe_allow_html=True,
        )

        st.markdown(
            '<div class="stat-row">'
            f'<div class="stat-card"><div class="num">{stats["total_eq"]}</div>'
            '<div class="lbl">方程式總數</div></div>'
            f'<div class="stat-card"><div class="num">{stats["inline"]}</div>'
            '<div class="lbl">行內公式</div></div>'
            f'<div class="stat-card"><div class="num">{stats["display"]}</div>'
            '<div class="lbl">展示公式</div></div>'
            f'<div class="stat-card"><div class="num">{stats["tables"]}</div>'
            '<div class="lbl">表格</div></div>'
            f'<div class="stat-card"><div class="num">{stats["headings"]}</div>'
            '<div class="lbl">標題</div></div>'
            "</div>",
            unsafe_allow_html=True,
        )

        tab_preview, tab_equations = st.tabs(["📄 Markdown 預覽", "🧮 方程式清單"])

        with tab_preview:
            st.markdown(text)

        with tab_equations:
            if equations:
                for idx, (eq_type, eq_latex) in enumerate(equations, 1):
                    tag_color = (
                        "#0d9488" if eq_type == "展示"
                        else "#6366f1" if eq_type == "行內"
                        else "#f59e0b"
                    )
                    st.markdown(
                        f"**#{idx}** &nbsp;"
                        f"<span style='background:{tag_color};color:#fff;"
                        f"padding:1px 8px;border-radius:10px;font-size:.75rem'>"
                        f"{eq_type}</span>",
                        unsafe_allow_html=True,
                    )
                    try:
                        st.latex(eq_latex)
                    except Exception:
                        st.code(eq_latex)
            else:
                st.info("未偵測到數學方程式。")

        # ── Step 3：匯出 ──
        st.markdown(
            '<div class="step-label"><span class="badge">3</span>產生並下載文件</div>',
            unsafe_allow_html=True,
        )

        if st.button("🚀 產生 Word 文件", type="primary", use_container_width=True):
            with st.spinner("轉換中，請稍候…"):
                if "MathType" in mode and "傳統" not in mode:
                    mmode = "latex"
                else:
                    mmode = "omml"

                doc = build_docx(
                    segments,
                    math_mode=mmode,
                    font_name=sel_font,
                    body_size=sel_size,
                    line_spacing=sel_spacing,
                    tcolor=sel_tcolor_hex,
                )
                buf = BytesIO()
                doc.save(buf)
                buf.seek(0)

            st.success("✅ 轉換完成！")

            if "OMML" in mode and "傳統" not in mode:
                st.download_button(
                    "📥 下載 Word（OMML 內建方程式）",
                    data=buf,
                    file_name="mathdoc_omml.docx",
                    mime="application/vnd.openxmlformats-officedocument"
                         ".wordprocessingml.document",
                    use_container_width=True,
                )

            elif "⭐" in mode:
                st.download_button(
                    "📥 下載 Word（MathType LaTeX 模式）",
                    data=buf,
                    file_name="mathdoc_mathtype.docx",
                    mime="application/vnd.openxmlformats-officedocument"
                         ".wordprocessingml.document",
                    use_container_width=True,
                )
                st.info(
                    "💡 **下一步：** Word 開啟 → MathType → Convert Equations → "
                    "勾選 **MathType translator text equations** → "
                    "Translator 選 **TeX -- AMS-LaTeX** → Convert\n\n"
                    "也可以選取單個 `$...$` 按 **Alt+\\\\** 即時轉換。"
                )

            elif "傳統" in mode:
                c1, c2 = st.columns(2)
                with c1:
                    st.download_button(
                        "📥 Word (OMML)", data=buf,
                        file_name="mathdoc_omml.docx",
                        mime="application/vnd.openxmlformats-officedocument"
                             ".wordprocessingml.document",
                        use_container_width=True,
                    )
                with c2:
                    html = build_mathml_html(segments)
                    st.download_button(
                        "📥 MathML 網頁檔",
                        data=html.encode("utf-8"),
                        file_name="equations_mathml.html",
                        mime="text/html",
                        use_container_width=True,
                    )
                st.info(
                    "💡 MathType → Convert Equations → Whole Document "
                    "→ MathType equations (OLE) → Convert"
                )
                with st.expander("📋 VBA 巨集（進階）"):
                    st.code(VBA_OMML, language="vb")

    elif not text.strip():
        st.markdown(
            '<div style="text-align:center;padding:3rem 0;color:#9ca3af">'
            '<p style="font-size:2.5rem;margin:0">📄</p>'
            "<p>在上方輸入框貼上 AI 產出的內容，即可開始轉換</p>"
            "</div>",
            unsafe_allow_html=True,
        )


if __name__ == "__main__":
    main()